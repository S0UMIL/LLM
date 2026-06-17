import streamlit as st
from langchain_ollama import OllamaLLM
from langchain_core.prompts import PromptTemplate
from docx import Document


st.title("Design Project Intake Form")

with st.form("intake_form"):
    project_name = st.text_input("Project Name")
    client_name = st.text_input("Client Name")
    domain = st.selectbox("Project Domain", ["Automation/Robotics", "Software/IT", "New Technology/R&D", "Industrial/Manufacturing"])
    project_purpose= st.selectbox("purpose",["Redesign","launching a new product","increasing sales"])
    target_audience = st.selectbox("target Audience",["daily consumers","office professionals","In company use"])
    client_budget = st.number_input("Enter your Budget")
    project_date = st.date_input("Enter the day by which you want your project to be ready")
    project_description = st.text_area("Project Description")
    project_type = st.selectbox("Project Type", ["Residential", "Commercial", "Other"])
    submitted = st.form_submit_button("Submit")
llm=OllamaLLM(model="llama3.2",temperature=0.7)

first_input=PromptTemplate(
    input_variables=['project_name','domain','budget','deadline','description'],
    template="""You are an assistant at an Tech Firm responsible for generating a text on project design report,based on the project's given {description}
    the project's name is {project_name}. the budget of the client is this {budget},keep in mind the solution you are gonna generate in text should be with respect to the {domain} strictly not halluicnating
    the deadline of the project is {deadline} make sure the solution is feasible with respect to date. start your answer with content of the text like an book index then jump into detail"""
)
chain1= first_input | llm
doc=Document()




if submitted:
    st.write("Form submitted!")
    st.write(f"Project: {project_name}, Client: {client_name}")
    st.write("generating response please wait.")

    generated_answer = chain1.invoke({
    "project_name": project_name,
    "domain": domain,
    "budget": client_budget,
    "deadline": project_date,
    "description": project_description
})
    st.write(generated_answer)
    doc.add_heading(project_name,level=1)
    for para in generated_answer.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.save("project_report.docx")
